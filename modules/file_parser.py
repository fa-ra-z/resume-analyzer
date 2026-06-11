"""
Universal Resume File Parser — MULTI-PAGE OPTIMIZED
─────────────────────────────────────────────────────
Accepts ANY common resume format and returns clean, AI-ready text
from ALL pages (handles 1, 2, 3+ page resumes flawlessly).

Supported formats
─────────────────
✓ PDF (text-based, any number of pages)    .pdf
✓ PDF (scanned / image-based, multi-page)  .pdf  → OCR fallback per page
✓ Microsoft Word (modern)                  .docx
✓ Microsoft Word (legacy)                  .doc
✓ Rich Text Format                         .rtf
✓ Plain Text                               .txt
✓ Images (resumes as photos)               .png .jpg .jpeg .webp .bmp .tiff

Features
────────
✓ Reads EVERY page (no page limit)
✓ Auto file-type detection (magic bytes + extension)
✓ Multi-column PDF layout detection per page
✓ Per-page OCR fallback for scanned PDFs
✓ Hyperlink extraction from all pages
✓ Ligature, smart-quote, and bullet normalization
✓ Page-number / header / footer stripping
✓ Image preprocessing for better OCR accuracy
✓ Cross-page paragraph rejoining
✓ Optional rich metadata extraction
"""

import io
import re
import unicodedata
from typing import Union, BinaryIO

import fitz   # PyMuPDF


# ─────────────────────────────────────────────────────────────
# OPTIONAL DEPENDENCIES — loaded only when needed
# ─────────────────────────────────────────────────────────────
def _try_import(module_name: str):
    try:
        return __import__(module_name)
    except ImportError:
        return None


# ─────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────
COMMON_LIGATURES = {
    "\ufb00": "ff",  "\ufb01": "fi",  "\ufb02": "fl",
    "\ufb03": "ffi", "\ufb04": "ffl", "\ufb05": "ft",
    "\ufb06": "st",
}

BULLET_CHARS = [
    "\u2022", "\u2023", "\u25E6", "\u2043", "\u2219",
    "\u25AA", "\u25AB", "\u25CB", "\u25CF", "\u25A0",
    "\u00B7", "\u2027", "\u2218", "\u29BE", "\u29BF",
    "►", "▶", "➤", "➢", "❖", "✦", "✧", "✪", "✶", "✱",
    "◆", "◇", "◉", "○", "●", "■", "□", "▪", "▫",
]

SMART_REPLACEMENTS = {
    "\u2018": "'",  "\u2019": "'",
    "\u201C": '"',  "\u201D": '"',
    "\u2013": "-",  "\u2014": "-",
    "\u2212": "-",  "\u2010": "-",
    "\u00A0": " ",  "\u202F": " ",  "\u2009": " ",
    "\u2026": "...",
    "\u00AD": "",
}

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".rtf", ".txt",
    ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}

MAGIC_BYTES = {
    b"%PDF":               "pdf",
    b"PK\x03\x04":         "docx",
    b"\xD0\xCF\x11\xE0":   "doc",
    b"{\\rtf":             "rtf",
    b"\x89PNG":            "png",
    b"\xFF\xD8\xFF":       "jpg",
    b"RIFF":               "webp",
    b"GIF8":               "gif",
    b"BM":                 "bmp",
    b"II*\x00":            "tiff",
    b"MM\x00*":            "tiff",
}


# ═════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ═════════════════════════════════════════════════════════════
def extract_text_from_pdf(uploaded_file, return_metadata: bool = False):
    """Backward-compatible alias — now accepts ALL file types."""
    return extract_text(uploaded_file, return_metadata=return_metadata)


def extract_text(
    uploaded_file: Union[BinaryIO, bytes],
    return_metadata: bool = False,
):
    """Universal text extractor — reads ALL pages of any resume format."""
    if hasattr(uploaded_file, "read"):
        raw_bytes = uploaded_file.read()
        filename = getattr(uploaded_file, "name", "") or ""
    elif isinstance(uploaded_file, (bytes, bytearray)):
        raw_bytes = bytes(uploaded_file)
        filename = ""
    else:
        raise ValueError("uploaded_file must be a file-like object or bytes")

    if not raw_bytes:
        return _empty_result(return_metadata, "Empty file")

    file_type = _detect_file_type(raw_bytes, filename)

    extracted_text = ""
    links_found = set()
    page_count = 0

    try:
        if file_type == "pdf":
            extracted_text, links_found, page_count = _extract_pdf(raw_bytes)

        elif file_type == "docx":
            extracted_text, links_found = _extract_docx(raw_bytes)

        elif file_type == "doc":
            extracted_text = _extract_doc(raw_bytes)

        elif file_type == "rtf":
            extracted_text = _extract_rtf(raw_bytes)

        elif file_type == "txt":
            extracted_text = _extract_txt(raw_bytes)

        elif file_type in ("png", "jpg", "jpeg", "webp", "bmp", "tiff", "gif"):
            extracted_text = _extract_image(raw_bytes)

        else:
            return _empty_result(
                return_metadata,
                f"Unsupported file type: {file_type or 'unknown'}"
            )

    except Exception as e:
        return _empty_result(return_metadata, f"Extraction failed: {str(e)}")

    extracted_text = _inject_missing_links(extracted_text, links_found)
    clean = _clean_text(extracted_text)

    if return_metadata:
        return {
            "text":         clean,
            "file_type":    file_type,
            "page_count":   page_count,
            "word_count":   len(clean.split()),
            "char_count":   len(clean),
            "links_found":  sorted(links_found),
            "emails":       _find_emails(clean),
            "phones":       _find_phones(clean),
            "linkedin":     _find_linkedin(clean, links_found),
            "github":       _find_github(clean, links_found),
            "error":        None,
        }
    return clean


# ═════════════════════════════════════════════════════════════
# FILE TYPE DETECTION
# ═════════════════════════════════════════════════════════════
def _detect_file_type(raw_bytes: bytes, filename: str) -> str:
    header = raw_bytes[:8]
    for magic, ftype in MAGIC_BYTES.items():
        if header.startswith(magic):
            if magic == b"RIFF" and b"WEBP" in raw_bytes[:16]:
                return "webp"
            elif magic == b"RIFF":
                continue
            return ftype

    if filename:
        ext = "." + filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if ext in SUPPORTED_EXTENSIONS:
            return ext.lstrip(".")

    try:
        raw_bytes[:1000].decode("utf-8")
        return "txt"
    except UnicodeDecodeError:
        pass

    return ""


# ═════════════════════════════════════════════════════════════
# PDF EXTRACTION — MULTI-PAGE OPTIMIZED
# ═════════════════════════════════════════════════════════════
def _extract_pdf(raw_bytes: bytes) -> tuple:
    """
    Extract text from EVERY page of the PDF.
    Auto-detects scanned pages and applies OCR per page.
    """
    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    pages = []
    all_links = set()
    pages_needing_ocr = []

    # ── First pass: extract text from every page ─────────
    for page_num, page in enumerate(doc):
        try:
            page_text = _extract_pdf_page(page)
            page_text = _strip_header_footer(page_text)

            # If this individual page has almost no text → mark for OCR
            if len(page_text.strip()) < 30:
                pages_needing_ocr.append(page_num)
                page_text = ""   # placeholder

            pages.append(page_text)

            # Extract hyperlinks from this page
            for link in page.get_links():
                uri = link.get("uri", "")
                if uri:
                    all_links.add(uri.strip())

        except Exception:
            pages.append("")
            pages_needing_ocr.append(page_num)

    page_count = len(pages)

    # ── Second pass: OCR any pages that returned no text ─
    if pages_needing_ocr:
        ocr_results = _ocr_specific_pages(doc, pages_needing_ocr)
        for page_num, ocr_text in ocr_results.items():
            if ocr_text and len(ocr_text.strip()) > len(pages[page_num].strip()):
                pages[page_num] = _strip_header_footer(ocr_text)

    # ── If ENTIRE document is empty → full OCR fallback ──
    combined = "\n\n".join(pages).strip()
    if len(combined) < 100:
        full_ocr = _ocr_pdf_pages(doc)
        if full_ocr and len(full_ocr) > len(combined):
            combined = full_ocr

    # ── Smart cross-page joining ─────────────────────────
    final_text = _smart_join_pages(pages)
    if not final_text.strip():
        final_text = combined

    doc.close()
    return final_text, all_links, page_count


def _extract_pdf_page(page) -> str:
    """Extract one PDF page in correct reading order (multi-column aware)."""
    try:
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_LIGATURES)["blocks"]
    except Exception:
        return page.get_text("text")

    items = []
    for block in blocks:
        if block.get("type", 0) != 0:
            continue
        for line in block.get("lines", []):
            line_text = "".join(span.get("text", "") for span in line.get("spans", []))
            line_text = line_text.strip()
            if line_text:
                items.append({
                    "x": line["bbox"][0],
                    "y": line["bbox"][1],
                    "text": line_text,
                })

    if not items:
        return page.get_text("text")

    # Detect two-column layout
    page_width = page.rect.width
    midpoint = page_width / 2
    left_count = sum(1 for it in items if it["x"] < midpoint - 20)
    right_count = sum(1 for it in items if it["x"] > midpoint + 20)
    is_two_column = (
        left_count > 5 and right_count > 5
        and abs(left_count - right_count) < max(left_count, right_count) * 0.6
    )

    if is_two_column:
        left  = sorted([it for it in items if it["x"] < midpoint],
                       key=lambda i: (round(i["y"], 1), i["x"]))
        right = sorted([it for it in items if it["x"] >= midpoint],
                       key=lambda i: (round(i["y"], 1), i["x"]))
        ordered = left + right
    else:
        ordered = sorted(items, key=lambda i: (round(i["y"], 1), i["x"]))

    lines, prev_y = [], None
    for it in ordered:
        if prev_y is not None and abs(it["y"] - prev_y) > 14:
            lines.append("")
        lines.append(it["text"])
        prev_y = it["y"]
    return "\n".join(lines)


def _smart_join_pages(pages: list) -> str:
    """
    Join pages intelligently — if a page ends mid-sentence and the next
    page starts with lowercase, glue them. Otherwise keep paragraph break.
    """
    if not pages:
        return ""

    result = pages[0].rstrip()
    for next_page in pages[1:]:
        next_page = next_page.strip()
        if not next_page:
            continue

        prev_ends_mid_sentence = (
            result and result[-1] not in ".!?:;)\"]'"
            and not result.endswith(("\n\n",))
        )
        next_starts_lower = next_page[0].islower() if next_page else False

        if prev_ends_mid_sentence and next_starts_lower:
            # Continuation — join with space
            result += " " + next_page
        else:
            # New section — paragraph break
            result += "\n\n" + next_page

    return result


def _ocr_specific_pages(doc, page_nums: list) -> dict:
    """Run OCR only on specific pages that need it (efficient)."""
    pytesseract = _try_import("pytesseract")
    PIL = _try_import("PIL")
    if not pytesseract or not PIL:
        return {}

    from PIL import Image
    results = {}
    for page_num in page_nums:
        try:
            page = doc[page_num]
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = _preprocess_image_for_ocr(img)
            text = pytesseract.image_to_string(img, lang="eng")
            results[page_num] = text
        except Exception:
            continue
    return results


def _ocr_pdf_pages(doc) -> str:
    """OCR every page (used when full document is empty)."""
    pytesseract = _try_import("pytesseract")
    PIL = _try_import("PIL")
    if not pytesseract or not PIL:
        return ""

    from PIL import Image
    pages_text = []
    for page in doc:
        try:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = _preprocess_image_for_ocr(img)
            text = pytesseract.image_to_string(img, lang="eng")
            pages_text.append(text)
        except Exception:
            continue
    return "\n\n".join(pages_text)


# ═════════════════════════════════════════════════════════════
# DOCX
# ═════════════════════════════════════════════════════════════
def _extract_docx(raw_bytes: bytes) -> tuple:
    docx = _try_import("docx")
    if not docx:
        return _extract_docx_via_docx2txt(raw_bytes), set()

    from docx import Document
    doc = Document(io.BytesIO(raw_bytes))

    parts = []
    links = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    try:
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype.lower():
                target = rel.target_ref
                if target and target.startswith("http"):
                    links.add(target)
    except Exception:
        pass

    return "\n".join(parts), links


def _extract_docx_via_docx2txt(raw_bytes: bytes) -> str:
    docx2txt = _try_import("docx2txt")
    if not docx2txt:
        return ""
    try:
        return docx2txt.process(io.BytesIO(raw_bytes)) or ""
    except Exception:
        return ""


# ═════════════════════════════════════════════════════════════
# LEGACY DOC
# ═════════════════════════════════════════════════════════════
def _extract_doc(raw_bytes: bytes) -> str:
    docx2txt = _try_import("docx2txt")
    if docx2txt:
        try:
            text = docx2txt.process(io.BytesIO(raw_bytes))
            if text and len(text.strip()) > 30:
                return text
        except Exception:
            pass

    textract = _try_import("textract")
    if textract:
        try:
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
                tmp.write(raw_bytes)
                tmp_path = tmp.name
            text = textract.process(tmp_path).decode("utf-8", errors="ignore")
            os.unlink(tmp_path)
            return text
        except Exception:
            pass

    return _extract_doc_fallback(raw_bytes)


def _extract_doc_fallback(raw_bytes: bytes) -> str:
    try:
        text = raw_bytes.decode("utf-8", errors="ignore")
        text = re.sub(r"[^\x20-\x7E\n\r\t]+", " ", text)
        text = re.sub(r"\s{3,}", "\n", text)
        return text if len(text.strip()) > 50 else ""
    except Exception:
        return ""


# ═════════════════════════════════════════════════════════════
# RTF / TXT / IMAGE
# ═════════════════════════════════════════════════════════════
def _extract_rtf(raw_bytes: bytes) -> str:
    striprtf = _try_import("striprtf")
    if striprtf:
        try:
            from striprtf.striprtf import rtf_to_text
            text = raw_bytes.decode("utf-8", errors="ignore")
            return rtf_to_text(text)
        except Exception:
            pass
    try:
        text = raw_bytes.decode("utf-8", errors="ignore")
        text = re.sub(r"\\[a-z]+\d* ?", "", text)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"\\\'[0-9a-f]{2}", "", text)
        return text
    except Exception:
        return ""


def _extract_txt(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore")


def _extract_image(raw_bytes: bytes) -> str:
    pytesseract = _try_import("pytesseract")
    PIL = _try_import("PIL")
    if not pytesseract or not PIL:
        return ""

    from PIL import Image
    try:
        img = Image.open(io.BytesIO(raw_bytes))
        img = _preprocess_image_for_ocr(img)
        return pytesseract.image_to_string(img, lang="eng")
    except Exception:
        return ""


def _preprocess_image_for_ocr(img):
    from PIL import Image, ImageOps, ImageFilter
    img = img.convert("L")
    img = ImageOps.autocontrast(img, cutoff=2)
    if img.width < 1500:
        scale = 1500 / img.width
        new_size = (int(img.width * scale), int(img.height * scale))
        img = img.resize(new_size, Image.LANCZOS)
    img = img.filter(ImageFilter.SHARPEN)
    return img


# ═════════════════════════════════════════════════════════════
# SHARED HELPERS
# ═════════════════════════════════════════════════════════════
def _empty_result(return_metadata: bool, error_msg: str):
    if return_metadata:
        return {
            "text": "", "file_type": "", "page_count": 0,
            "word_count": 0, "char_count": 0,
            "links_found": [], "emails": [], "phones": [],
            "linkedin": "", "github": "", "error": error_msg,
        }
    return ""


def _strip_header_footer(text: str) -> str:
    cleaned = []
    for line in text.split("\n"):
        s = line.strip()
        if re.fullmatch(r"\d{1,3}", s):                          continue
        if re.fullmatch(r"(?i)page\s*\d+(\s*of\s*\d+)?", s):     continue
        if re.fullmatch(r"\d+\s*/\s*\d+", s):                    continue
        cleaned.append(line)
    return "\n".join(cleaned)


def _inject_missing_links(text: str, links: set) -> str:
    if not links:
        return text
    text_lower = text.lower()
    additions = []
    for link in links:
        link_lower = link.lower()
        if link_lower.startswith(("mailto:", "tel:")):
            continue
        slug = link.rstrip("/").split("/")[-1].lower()
        if link_lower not in text_lower and slug not in text_lower:
            additions.append(link)
    if additions:
        text += "\n\n" + "\n".join(additions)
    return text


def _clean_text(text: str) -> str:
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    for lig, repl in COMMON_LIGATURES.items():
        text = text.replace(lig, repl)
    for smart, ascii_char in SMART_REPLACEMENTS.items():
        text = text.replace(smart, ascii_char)
    for b in BULLET_CHARS:
        text = text.replace(b, "•")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"•\s*", "• ", text)
    text = "".join(
        ch for ch in text
        if ch == "\n" or ch == "\t" or (ord(ch) >= 32 and ord(ch) != 127)
    )
    return text.strip()


def _find_emails(text: str) -> list:
    pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    return list(dict.fromkeys(re.findall(pattern, text)))


def _find_phones(text: str) -> list:
    pattern = r"(\+?\d{1,3}[\s-]?)?\(?\d{3,5}\)?[\s.-]?\d{3,5}[\s.-]?\d{3,5}"
    phones = []
    for m in re.finditer(pattern, text):
        raw = m.group(0).strip()
        digits = re.sub(r"\D", "", raw)
        if 10 <= len(digits) <= 15:
            phones.append(raw)
    return list(dict.fromkeys(phones))


def _find_linkedin(text: str, links: set) -> str:
    for link in links:
        if "linkedin.com" in link.lower():
            return link.strip()
    m = re.search(r"https?://(?:www\.)?linkedin\.com/[^\s)\]]+", text, re.I)
    if m: return m.group(0).rstrip(".,;")
    m = re.search(r"linkedin\.com/in/[A-Za-z0-9_-]+", text, re.I)
    if m: return "https://" + m.group(0)
    return ""


def _find_github(text: str, links: set) -> str:
    for link in links:
        if "github.com" in link.lower():
            return link.strip()
    m = re.search(r"https?://(?:www\.)?github\.com/[^\s)\]]+", text, re.I)
    if m: return m.group(0).rstrip(".,;")
    m = re.search(r"github\.com/[A-Za-z0-9_-]+", text, re.I)
    if m: return "https://" + m.group(0)
    return ""